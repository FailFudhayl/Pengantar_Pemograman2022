from docx import Document
from docx.shared import Inches

word = Document()
word.add_heading('Data Diri Fail')

word.add_paragraph("Nama : Andi Ahmad Fail Fudhayl")
word.add_paragraph("NIM : H071221026")
word.add_paragraph("Tanggal Lahir : 10 Januari 2004")

word.save('Fail.docx')
